import io
import openpyxl
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

LEVELS = ["初級", "中級", "高級", "研經"]
PERIODS = ["日", "夜"]
FONT_NAME = "微軟正黑體"


def detect_level(filename: str):
    for lv in LEVELS:
        if lv in filename:
            return lv
    return None


def detect_period(filename: str):
    """
    偵測檔名中的日班／夜班，回傳 '日' 或 '夜'；偵測不到回傳 None。

    支援兩種常見排法：
    1. 班級關鍵字「後面」緊接著日／夜，例如「初級夜」
    2. 班級關鍵字「前面」緊接著日／夜，例如「四日研經」
       （「日」代表日班，前面的數字代表星期幾，例如「四」= 星期四）

    找不到以上兩種明確排法時，才退而求其次找「日班」／「夜班」這種完整字樣。
    這樣可以避免誤判像單純「四日禪修營」這種與日夜班完全無關、卻剛好含有「日」字的檔名。
    """
    level = detect_level(filename)
    if level:
        idx = filename.find(level)
        end_idx = idx + len(level)
        if end_idx < len(filename) and filename[end_idx] in ("日", "夜"):
            return filename[end_idx]
        if idx > 0 and filename[idx - 1] in ("日", "夜"):
            return filename[idx - 1]
    if "夜班" in filename:
        return "夜"
    if "日班" in filename:
        return "日"
    return None


def load_course_lookup(file_bytes, filename="course.xlsx"):
    """Returns {level_sheet_name: {mm/dd: course_name}}"""
    wb = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=True)
    lookup = {}
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        date_to_course = {}
        for row in ws.iter_rows(min_row=3, values_only=False):
            date_cell = row[0].value
            course_cell = row[2].value if len(row) > 2 else None
            if date_cell is not None and hasattr(date_cell, "strftime"):
                mmdd = date_cell.strftime("%m/%d")
                date_to_course[mmdd] = course_cell if course_cell else "(課程名稱未提供)"
        lookup[sheet_name] = date_to_course
    return lookup


KNOWN_CODES = {"V", "O", "L", "LL", "M", "ML"}
ABSENCE_CODE = "O"


def _extract_leading_code(s):
    """從字串開頭抓出已知代碼，例如「V座禪」開頭是 V。找不到回傳 None。"""
    for code in sorted(KNOWN_CODES, key=len, reverse=True):
        if s.startswith(code):
            return code
    return None


def _classify_status_cell(raw_value):
    """
    Classifies one attendance-status cell.
    Returns (kind, detail):
      kind = "absence"     -> counts toward makeup notice
      kind = "present"     -> recognized code, not an absence
      kind = "ambiguous"   -> multiple codes combined in one cell (e.g. "O / V")
      kind = "unrecognized"-> non-empty but not a known code
      kind = "empty"       -> no value recorded
    """
    if raw_value is None:
        return "empty", None
    s = str(raw_value).strip()
    if not s:
        return "empty", None

    if "/" in s:
        parts = [p.strip() for p in s.split("/") if p.strip()]
        if len(parts) > 1 and all(p in KNOWN_CODES for p in parts):
            return "ambiguous", s

    if s in KNOWN_CODES:
        code = s
    else:
        code = None
        tokens = s.split()
        if tokens:
            # 「標籤 代碼」格式，例如「夜研 V」→ 代碼在最後一個字
            if tokens[-1] in KNOWN_CODES:
                code = tokens[-1]
            # 「代碼 (備註)」格式，例如「V (8/20補M)」「M (7/30補)」→ 代碼在第一個字
            elif tokens[0] in KNOWN_CODES:
                code = tokens[0]
        if code is None:
            # 代碼與中文字直接相連、沒有空白，例如「V座禪」
            code = _extract_leading_code(s)

    if code == ABSENCE_CODE:
        return "absence", s
    if code in KNOWN_CODES:
        return "present", s
    return "unrecognized", s


def load_attendance(file_bytes, filename):
    """Returns (class_title, level, [student dicts], warnings) for one attendance file."""
    wb = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=True)
    ws = wb.active
    title = ws["A1"].value or filename
    title_str = str(title) if title else ""

    # 優先用檔名判斷班級／日夜；檔名判斷不出來時（例如檔名被簡化、縮寫過），
    # 改抓 Excel 內 A1 儲存格的標題文字來判斷，因為標題通常保留完整班級名稱。
    level = detect_level(filename) or detect_level(title_str)
    period = detect_period(filename) or detect_period(title_str)

    # date headers live in row 3, columns E..N (index 4..13, 0-based)
    header_row = next(ws.iter_rows(min_row=3, max_row=3, values_only=True))
    raw_date_headers = header_row[4:14]
    date_headers = []
    for h in raw_date_headers:
        if hasattr(h, "strftime"):
            date_headers.append(h.strftime("%m/%d"))
        elif h is None:
            date_headers.append("?")
        else:
            date_headers.append(str(h).strip())

    students = []
    warnings = []
    for row in ws.iter_rows(min_row=4, values_only=False):
        if len(row) < 2 or row[1].value is None:
            continue
        no_cell, name_cell, dharma_cell, group_cell = row[0], row[1], row[2], row[3]
        status_cells = row[4:14]

        absence_dates = []
        for col_idx, cell in enumerate(status_cells):
            mmdd = date_headers[col_idx] if col_idx < len(date_headers) else "?"
            kind, detail = _classify_status_cell(cell.value)
            if kind == "absence":
                absence_dates.append(mmdd)
            elif kind == "ambiguous":
                warnings.append(
                    f"{name_cell.value}（{filename}）於 {mmdd} 的簽到狀態為「{detail}」，"
                    f"格式模糊（同一格出現多種代碼），請人工確認是否算缺曠。"
                )
            elif kind == "unrecognized":
                warnings.append(
                    f"{name_cell.value}（{filename}）於 {mmdd} 的簽到狀態為「{detail}」，"
                    f"不是可辨識的代碼（V/O/L/LL/M/ML），請人工確認。"
                )

        if not absence_dates:
            continue

        students.append({
            "no": no_cell.value,
            "name": name_cell.value,
            "dharma_name": dharma_cell.value,
            "group": group_cell.value,
            "absence_count": len(absence_dates),
            "absence_dates": absence_dates,
            "level": level,
            "period": period,
            "class_title": title,
            "source_file": filename,
        })
    return title, level, period, students, warnings


PERIOD_LABELS = {"日": "日班", "夜": "夜班"}

# 研經班若有分開的日/夜課表（分頁「研經日」「研經夜」），會優先對應；
# 若課程名稱 Excel 沒有分開（只有單一「研經」分頁），則自動退回共用該分頁。
LEVELS_WITH_SEPARATE_PERIOD_SCHEDULE = {"研經"}


def resolve_course_sheet(level, period, course_lookup):
    """
    決定要用課程名稱 Excel 裡的哪個分頁來查課表。
    優先找「班級+日夜」專屬分頁（例如「研經日」），
    找不到的話自動退回只用班級名稱的分頁（例如「研經」）。
    回傳 (date_map, matched_key)；matched_key 為 None 代表完全找不到對應分頁。
    """
    candidates = []
    if level and period:
        candidates.append(f"{level}{period}")
    if level:
        candidates.append(level)
    for key in candidates:
        if key in course_lookup:
            return course_lookup[key], key
    return {}, None


def format_level_period(level, period):
    """組合班別顯示文字，例如「初級夜班」「研經日班」。"""
    level_str = level or ""
    period_label = PERIOD_LABELS.get(period, "")
    return f"{level_str}{period_label}"


def build_student_records(attendance_files, course_lookup):
    """
    attendance_files: list of (filename, bytes)
    course_lookup: dict from load_course_lookup
    Returns list of student dicts, each with resolved 'items': [{date, course}]
    """
    all_students = []
    warnings = []
    for filename, file_bytes in attendance_files:
        title, level, period, students, cell_warnings = load_attendance(file_bytes, filename)
        warnings.extend(cell_warnings)
        if level is None:
            warnings.append(f"檔名「{filename}」無法判斷班級（初級/中級/高級/研經），將無法對應課程名稱。")

        date_map, matched_key = resolve_course_sheet(level, period, course_lookup)
        if matched_key is None and level is not None:
            warnings.append(
                f"檔名「{filename}」在課程名稱 Excel 中找不到「{level}」對應的分頁，"
                f"該班所有缺曠日期都無法對應課程，請確認課程名稱 Excel 的分頁名稱。"
            )
        elif period is None and level in LEVELS_WITH_SEPARATE_PERIOD_SCHEDULE:
            has_split_schedule = any(
                f"{level}{p}" in course_lookup for p in PERIOD_LABELS
            )
            if has_split_schedule:
                warnings.append(
                    f"檔名「{filename}」屬於{level}班但無法判斷日夜，"
                    f"由於課程名稱 Excel 裡有分開的{level}日／{level}夜課表，可能對應到錯誤的課表，請人工確認。"
                )
        for s in students:
            items = []
            for d in s["absence_dates"]:
                course = date_map.get(d)
                if course is None:
                    course = "(查無對應課程，請人工確認)"
                    warnings.append(f"{s['name']}（{filename}）的缺曠日期 {d} 在課程名稱中找不到對應課程。")
                items.append({"date": d, "course": course})
            s["items"] = items
        all_students.extend(students)
    return all_students, warnings


# ---------------- docx generation (python-docx) ----------------

def _set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "999999")
        borders.append(el)
    tblPr.append(borders)


def _set_cell_shading(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tcPr.append(shd)


def _style_run(run, size=11, bold=False):
    run.font.name = FONT_NAME
    run.font.size = Pt(size)
    run.font.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), FONT_NAME)


def _add_paragraph(doc, text="", size=11, bold=False, align=None, space_after=6):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    _style_run(run, size=size, bold=bold)
    return p


TEMPLATE_FONT = "標楷體"
TITLE_SIZE = 18
BODY_SIZE = 14
ORG_NAME = "普印精舍"

DEFAULT_DEADLINE = "9月2日"
DEFAULT_HOURS_NOTE = "於周一至周日早上9:00~20:00至精舍櫃檯報到完成補課，"
DEFAULT_LATE_NOTE = "若超過補課期限，請事先預約補課時間，以利事前安排。"


def _tmpl_run(paragraph, text, size=BODY_SIZE, bold=True, underline=False):
    run = paragraph.add_run(text)
    run.font.name = TEMPLATE_FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.underline = underline
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), TEMPLATE_FONT)
    return run


def _add_text_with_underlined_phrases(paragraph, text, underline_phrases, size=BODY_SIZE):
    """
    Writes `text` into `paragraph`, automatically underlining the first
    occurrence of any phrase in `underline_phrases` that appears in it —
    so the original template's underline styling (e.g. under "事先預約")
    survives even if the surrounding wording is edited.
    """
    remaining = text
    offset = 0
    match_start, match_phrase = None, None
    for phrase in underline_phrases:
        if phrase and phrase in remaining:
            i = remaining.index(phrase)
            if match_start is None or i < match_start:
                match_start, match_phrase = i, phrase
    if match_phrase is None:
        _tmpl_run(paragraph, text, size=size)
        return
    before = remaining[:match_start]
    after = remaining[match_start + len(match_phrase):]
    if before:
        _tmpl_run(paragraph, before, size=size)
    _tmpl_run(paragraph, match_phrase, size=size, underline=True)
    if after:
        _tmpl_run(paragraph, after, size=size)


COMPACT_TITLE_SIZE = 13
COMPACT_BODY_SIZE = 10.5


def _write_notice_content(doc, student, org_name, deadline_text, hours_note, late_note,
                           title_size=TITLE_SIZE, body_size=BODY_SIZE, space_after=6,
                           blank_lines=True):
    """
    把一位學員的補課通知單內容寫進 `doc`。
    full 版（blank_lines=True）在各區塊間留空行，維持原本官方範本的排版；
    精簡版（blank_lines=False，合併多人版使用）省略空行、用較小字級，
    盡量在一頁內塞下多份完整通知單以節省紙張。
    除了最後一段之外，每一段都設定 keep_with_next，
    讓 Word 分頁時盡量把整張通知單留在同一頁，不會被硬生生從中間切開。
    """
    paragraphs = []

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(space_after)
    _tmpl_run(p, f"{org_name}–補課通知單", size=title_size)
    paragraphs.append(p)

    if blank_lines:
        paragraphs.append(doc.add_paragraph())

    level_display = format_level_period(student.get("level"), student.get("period"))
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    _tmpl_run(
        p,
        f"班別：{level_display}　　組別：{student.get('group') or ''}　　"
        f"姓名：{student.get('name','')}　　"
        f"法名：{student.get('dharma_name') or ''}",
        size=body_size,
    )
    paragraphs.append(p)

    if blank_lines:
        paragraphs.append(doc.add_paragraph())

    for item in student.get("items", []):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(space_after)
        _tmpl_run(p, f"缺課日期：{item['date']}　　缺課名稱：{item['course']}", size=body_size)
        paragraphs.append(p)

    if blank_lines:
        paragraphs.append(doc.add_paragraph())

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    _tmpl_run(p, "請在", size=body_size)
    _tmpl_run(p, deadline_text, size=body_size, underline=True)
    _tmpl_run(p, "前，", size=body_size)
    _add_text_with_underlined_phrases(p, hours_note, underline_phrases=[], size=body_size)
    paragraphs.append(p)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    _add_text_with_underlined_phrases(p, late_note, underline_phrases=["事先預約"], size=body_size)
    paragraphs.append(p)

    # 除了整份通知單的最後一段，其餘都設定 keep_with_next，
    # 讓這份通知單盡量整份留在同一頁，不會被分頁切斷。
    for p in paragraphs[:-1]:
        p.paragraph_format.keep_with_next = True

    return paragraphs


def _add_thick_divider(doc):
    """加一條粗分隔線，用來在合併文件裡區隔不同學員的通知單（方便列印後裁剪分發）。"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "24")  # 3pt 粗線
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def build_notice_docx(student, org_name=ORG_NAME, deadline_text=DEFAULT_DEADLINE,
                       hours_note=DEFAULT_HOURS_NOTE, late_note=DEFAULT_LATE_NOTE):
    """
    Build a single-student makeup-class notice matching the organization's
    official Word template (普印精舍-補課通知單).
    Blanks in the original hand-filled template are replaced with the
    student's actual data; the number of 缺課日期/缺課名稱 lines matches
    that student's actual absence count.
    """
    doc = Document()
    section = doc.sections[0]
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)

    _write_notice_content(
        doc, student, org_name, deadline_text, hours_note, late_note,
        title_size=TITLE_SIZE, body_size=BODY_SIZE, space_after=6, blank_lines=True,
    )

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def build_combined_notice_docx(students, org_name=ORG_NAME, deadline_text=DEFAULT_DEADLINE,
                                hours_note=DEFAULT_HOURS_NOTE, late_note=DEFAULT_LATE_NOTE):
    """
    把所有學員的補課通知單合併成同一份 Word 文件：
    版面精簡、字級縮小，盡量在同一頁塞下多份「完整」的通知單以節省紙張，
    每位學員的通知單之間用一條粗線分隔，方便列印後沿線裁剪分發給個人。
    """
    doc = Document()
    section = doc.sections[0]
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)

    for idx, student in enumerate(students):
        _write_notice_content(
            doc, student, org_name, deadline_text, hours_note, late_note,
            title_size=COMPACT_TITLE_SIZE, body_size=COMPACT_BODY_SIZE,
            space_after=2, blank_lines=False,
        )
        if idx < len(students) - 1:
            _add_thick_divider(doc)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
